import re
from docx import Document

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

def extract_text_from_docx(docx_path):
    """
    从DOCX文件中提取文本内容
    """
    try:
        doc = Document(docx_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(' | '.join(row_text))

        content = '\n'.join(full_text)

        content = re.sub(r'\n{3,}', '\n\n', content)

        return content

    except Exception as e:
        print(f"提取DOCX内容时出错: {str(e)}")
        raise

def extract_text_from_pdf(pdf_path):
    """
    从PDF文件中提取文本内容
    """
    if PdfReader is None:
        raise ImportError("PyPDF2库未安装，请运行: pip install PyPDF2")

    try:
        reader = PdfReader(pdf_path)
        full_text = []

        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                full_text.append(text.strip())

        content = '\n'.join(full_text)
        content = re.sub(r'\n{3,}', '\n\n', content)

        return content

    except Exception as e:
        print(f"提取PDF内容时出错: {str(e)}")
        raise

def extract_structured_content(docx_path):
    """
    提取结构化的文档内容（标题、段落、列表等）
    """
    try:
        doc = Document(docx_path)
        structured = {
            'title': '',
            'sections': [],
            'paragraphs': [],
            'tables': [],
            'lists': []
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if para.style.name.startswith('Heading'):
                structured['sections'].append({
                    'level': para.style.name,
                    'text': text
                })
            else:
                structured['paragraphs'].append(text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            structured['tables'].append(table_data)

        return structured

    except Exception as e:
        print(f"提取结构化内容时出错: {str(e)}")
        raise

def parse_requirements(content):
    """
    解析需求内容，提取关键信息
    """
    requirements = {
        'functions': [],
        'screens': [],
        'buttons': [],
        'inputs': [],
        'flows': []
    }

    function_keywords = ['功能', '模块', '页面', '系统', '管理', '设置']
    screen_keywords = ['页面', '界面', '窗口', '弹窗', '对话框']
    button_keywords = ['按钮', '点击', '提交', '保存', '取消', '删除', '编辑', '查询']
    input_keywords = ['输入', '填写', '编辑', '修改']

    for keyword in function_keywords:
        if keyword in content:
            requirements['functions'].append(keyword)

    for keyword in screen_keywords:
        if keyword in content:
            requirements['screens'].append(keyword)

    for keyword in button_keywords:
        if keyword in content:
            requirements['buttons'].append(keyword)

    for keyword in input_keywords:
        if keyword in content:
            requirements['inputs'].append(keyword)

    return requirements

if __name__ == '__main__':
    import sys

    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
        try:
            content = extract_text_from_docx(docx_path)
            print("=" * 50)
            print("文档内容提取成功:")
            print("=" * 50)
            print(content[:1000])
            if len(content) > 1000:
                print(f"\n... (共 {len(content)} 字符)")
        except Exception as e:
            print(f"提取失败: {str(e)}")
    else:
        print("用法: python docx_parser.py <docx文件路径>")
